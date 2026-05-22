"""DOCX text extraction (Phase 15).

Uses python-docx to walk paragraphs AND tables. Most .docx files in financial
consulting context (Олья's use case) have tables for client metrics — those
are the highest-value content. python-docx misses some edge cases (tracked
changes, embedded objects, comments) — for v1 we extract what's there and
flag what couldn't be read in `_meta`.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any


def extract_text(path: str, max_chars: int | None = None) -> dict:
    """Extract text from a .docx file. Returns {text, paragraphs_count,
    tables_count, chars, truncated, file_name}.

    paragraphs are emitted line-by-line; tables emitted with `|` cell separators
    so the agent sees structure. Section breaks between paragraphs and tables
    are marked with blank lines.

    Raises FileNotFoundError if path doesn't exist.
    Raises ValueError if not a .docx (does NOT support legacy .doc — different format).
    """
    import docx  # python-docx

    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if p.suffix.lower() != ".docx":
        raise ValueError(
            f"Not a .docx file (got {p.suffix!r}). Legacy .doc is unsupported — "
            "convert to .docx via LibreOffice or Word first."
        )

    doc = docx.Document(str(p))

    parts: list[str] = []
    para_count = 0
    table_count = 0

    # Walk top-level body in document order: paragraphs interleaved with tables.
    # python-docx exposes paragraphs/tables as separate lists, but we want order.
    # Pre-build element→object maps so each child lookup is O(1) instead of
    # O(N) scan — important for documents with 1000+ paragraphs.
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map = {t._element: t for t in doc.tables}
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1]  # strip namespace
        if tag == "p":
            para = para_map.get(child)
            if para is not None:
                text = para.text.strip()
                if text:
                    parts.append(text)
                para_count += 1
        elif tag == "tbl":
            tbl = tbl_map.get(child)
            if tbl is not None:
                table_count += 1
                parts.append("")  # blank line before table
                for row in tbl.rows:
                    row_cells = [_cell_text(cell) for cell in row.cells]
                    parts.append(" | ".join(row_cells))
                parts.append("")  # blank line after

    text = "\n".join(parts)
    truncated = False
    if max_chars is not None and len(text) > max_chars:
        text = text[:max_chars]
        truncated = True

    return {
        "text": text,
        "paragraphs_count": para_count,
        "tables_count": table_count,
        "chars": len(text),
        "truncated": truncated,
        "file_name": p.name,
    }


def _cell_text(cell: Any) -> str:
    """Extract text from a docx table cell. Cells can contain multiple paragraphs."""
    pieces = []
    for para in cell.paragraphs:
        t = para.text.strip()
        if t:
            pieces.append(t)
    return " ".join(pieces)
