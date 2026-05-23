"""DOCX template rendering.

Sibling of `_docx_extract.py`. Where extract turns DOCX into text, this
module fills a DOCX template with caller-supplied values and writes a
new file. Built specifically for the TrueStats weekly-report workflow:

    template:  «Выручка за неделю: {revenue}, кол-во заказов: {orders}»
    data:      {"revenue": "1.2M ₽", "orders": "342"}
    output:    «Выручка за неделю: 1.2M ₽, кол-во заказов: 342»

Placeholder syntax: `{var_name}` — same as Python f-strings, but
braces are interpreted as plain text by Word, so the template stays
editable in Word/LibreOffice without escape gymnastics.

Hard problem the implementation solves: in DOCX, a paragraph is a
sequence of formatted "runs". If a placeholder is split across runs
(`{var` in one, `_name}` in the next — happens when Word's autosave
re-flows text), a naïve text-level replace misses it. We rebuild the
paragraph text from its runs, do the substitution, write the result
back into the first run, and clear the rest. The first run's
formatting wins — fine for most templates where placeholders share
the surrounding paragraph's style.

Tables: every cell is itself a list of paragraphs; we walk them via
the same paragraph-rewrite helper.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

# Placeholder regex — `{var_name}`. Only word characters (letters,
# digits, underscore) inside the braces so {revenue.usd} etc. don't
# accidentally match formatting noise. Two-pass extraction (find +
# substitute) uses the same pattern so the contract is consistent.
_PLACEHOLDER_RE = re.compile(r"\{([A-Za-z_][A-Za-z0-9_]*)\}")


def _rewrite_paragraph(paragraph, data: dict, missing: set[str]) -> int:
    """Substitute placeholders in one Word paragraph. Returns count of
    replacements made. Mutates the paragraph's runs in place.

    Adds any referenced-but-not-supplied variable names to `missing`
    so the caller can report them. Missing variables are LEFT in place
    (template still readable) rather than being replaced with empty
    strings — easier to spot a forgotten field at review time."""
    if not paragraph.runs:
        return 0
    full_text = "".join(r.text for r in paragraph.runs)
    if not _PLACEHOLDER_RE.search(full_text):
        return 0

    replacements = [0]

    def _sub(match):
        name = match.group(1)
        if name in data:
            replacements[0] += 1
            return str(data[name])
        missing.add(name)
        return match.group(0)  # leave the placeholder intact

    new_text = _PLACEHOLDER_RE.sub(_sub, full_text)
    if new_text == full_text:
        return 0

    # Write rebuilt text into the first run; clear the rest. The first
    # run's formatting (bold/italic/font) wins — acceptable trade-off
    # given that templates usually keep placeholders in plain runs.
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return replacements[0]


def list_placeholders(template_path: str) -> dict:
    """Return every `{var}` referenced by the template. Useful for the
    agent to know what fields it needs to provide before calling render.

    Returns {ok, data: {placeholders: [...], paragraph_count, table_count},
             _meta}."""
    import docx  # python-docx

    p = Path(template_path)
    if not p.exists():
        return {"ok": False, "error_kind": "not_found",
                "error": f"template not found: {template_path}"}
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error_kind": "bad_input",
                "error": f"not a .docx file: {p.suffix!r}"}

    doc = docx.Document(str(p))
    found: set[str] = set()
    para_count = 0
    table_count = 0

    for para in doc.paragraphs:
        para_count += 1
        for m in _PLACEHOLDER_RE.finditer(para.text):
            found.add(m.group(1))

    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for m in _PLACEHOLDER_RE.finditer(para.text):
                        found.add(m.group(1))

    return {
        "ok": True,
        "data": {
            "placeholders": sorted(found),
            "paragraph_count": para_count,
            "table_count": table_count,
        },
        "_meta": {"template_path": str(p.resolve())},
    }


def render(template_path: str, output_path: str, data: dict) -> dict:
    """Fill the template with `data` and write to `output_path`.

    Missing keys are reported in the return envelope (`missing_vars`)
    but DO NOT raise — the placeholder is left intact in the output so
    the user can spot it on review. Extra keys in `data` that don't
    correspond to any placeholder are silently ignored.

    Returns {ok, data: {output_path, replacements_made, missing_vars},
             _meta}."""
    import docx  # python-docx

    p = Path(template_path)
    if not p.exists():
        return {"ok": False, "error_kind": "not_found",
                "error": f"template not found: {template_path}"}
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error_kind": "bad_input",
                "error": f"not a .docx file: {p.suffix!r}"}
    if not isinstance(data, dict):
        return {"ok": False, "error_kind": "bad_input",
                "error": f"data must be a dict, got {type(data).__name__}"}

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(str(p))
    missing: set[str] = set()
    replacements = 0

    for para in doc.paragraphs:
        replacements += _rewrite_paragraph(para, data, missing)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replacements += _rewrite_paragraph(para, data, missing)

    doc.save(str(out))
    return {
        "ok": True,
        "data": {
            "output_path": str(out.resolve()),
            "replacements_made": replacements,
            "missing_vars": sorted(missing),
        },
        "_meta": {"template_path": str(p.resolve())},
    }
