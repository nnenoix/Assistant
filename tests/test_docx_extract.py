"""Unit tests for src/tools/_docx_extract.py."""
from pathlib import Path

import pytest

from src.tools import _docx_extract


def _make_sample_docx(path: Path) -> None:
    """Create a sample .docx file with paragraphs and a table."""
    import docx
    doc = docx.Document()
    doc.add_paragraph("Это первый параграф.")
    doc.add_paragraph("Это второй параграф со словами.")

    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "Header A"
    tbl.cell(0, 1).text = "Header B"
    tbl.cell(0, 2).text = "Header C"
    tbl.cell(1, 0).text = "Cell A1"
    tbl.cell(1, 1).text = "Cell B1"
    tbl.cell(1, 2).text = "Cell C1"

    doc.add_paragraph("Финальный параграф после таблицы.")
    doc.save(str(path))


def test_extract_simple_docx(tmp_path):
    p = tmp_path / "sample.docx"
    _make_sample_docx(p)
    r = _docx_extract.extract_text(str(p))
    assert r["paragraphs_count"] == 3
    assert r["tables_count"] == 1
    assert "первый параграф" in r["text"]
    assert "Header A | Header B | Header C" in r["text"]
    assert "Cell A1 | Cell B1 | Cell C1" in r["text"]
    assert "Финальный параграф" in r["text"]
    assert r["chars"] == len(r["text"])
    assert r["truncated"] is False
    assert r["file_name"] == "sample.docx"


def test_extract_truncation(tmp_path):
    p = tmp_path / "sample.docx"
    _make_sample_docx(p)
    r = _docx_extract.extract_text(str(p), max_chars=50)
    assert r["chars"] == 50
    assert r["truncated"] is True


def test_extract_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        _docx_extract.extract_text("/nonexistent.docx")


def test_extract_rejects_non_docx(tmp_path):
    p = tmp_path / "fake.doc"
    p.write_text("not a real doc", encoding="utf-8")
    with pytest.raises(ValueError, match="Not a .docx"):
        _docx_extract.extract_text(str(p))
